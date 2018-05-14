# -*- coding: utf-8 -*-
"""
Created on Mon May 14 2018

@author: dang

Written for Python 2.7
"""

from __future__ import unicode_literals
import docx
import sys
reload(sys)
sys.setdefaultencoding('utf8')
import csv

def read_format(format_toggle_bolds, format_toggle_italics, default_bold, default_italic):
    word_format_bolds = [default_bold for item in format_toggle_bolds]
    word_format_italics = [default_italic for item in format_toggle_italics]
    for k in range(len(format_toggle_bolds)):
        # Using style name doesn't work with different languages, example:
        # English Word file: Bold, Not Italic, German: Fett,Nicht kursiv
        if format_toggle_bolds[k] == True:
            word_format_bolds[k] = not word_format_bolds[k]
    for k in range(len(format_toggle_italics)):
        if format_toggle_italics[k] == True:
            word_format_italics[k] = not word_format_italics[k]
    return word_format_bolds, word_format_italics

def split_capital(phrase):
    """Split capital and small words.
    Capital words are only allowed at the first half of the phrase.
    If a capital word is after a small word, it is NOT marked capital.
    """
    format_capital = []
    words = phrase.split()
    flag_prev_word_capital = True
    for k in range(len(words)):
        if words[k].isupper() and flag_prev_word_capital:
            format_capital.append(True)
        else:
            format_capital.append(False)
            flag_prev_word_capital = False
    return words, format_capital

def merge_similar_series(words, case_capital, format_bold, format_italic):
    newwords = list(words)
    newcase_capital = list(case_capital)
    newformat_bold = list(format_bold)
    newformat_italic = list(format_italic)
    for k in range(len(words)-1,0,-1):
        if newcase_capital[k] == newcase_capital[k-1] \
        and newformat_bold[k] == newformat_bold[k-1] \
        and newformat_italic[k] == newformat_italic[k-1]: 
            newwords[k-1] = ' '.join([newwords[k-1], newwords[k]])
            del newwords[k]
            del newcase_capital[k]
            del newformat_bold[k]
            del newformat_italic[k]
    return newwords, newcase_capital, newformat_bold, newformat_italic

def merge_with_comment_phrase(words, case_capital, format_bold, format_italic):
    """If the word is in (), it should be a comment for the previous one,
    hence be added to the previous one.
    """
    newwords = list(words)
    newcase_capital = list(case_capital)
    newformat_bold = list(format_bold)
    newformat_italic = list(format_italic)
    for k in range(len(words)-1,0,-1):
        if len(newwords[k].strip(','))>0 and newwords[k].strip(',')[0]=='(' and newwords[k].strip(',')[-1]==')':
            newwords[k-1] = ' '.join([newwords[k-1], newwords[k]])
            del newwords[k]
            del newcase_capital[k]
            del newformat_bold[k]
            del newformat_italic[k]
    return newwords, newcase_capital, newformat_bold, newformat_italic

def re_parse(word_texts, word_format_bolds, word_format_italics):
    """Fix the leftover mistake in the data parsed, some example items:
    THAN gate road
    TOÁN, S_CHÊ figure 
    hình 
    (bản vẽ) - is a separate item, in fact it is a comment for previous item
    AB-Betrieb m Đ_TỬ dass AB mode chê'độ hạng AB, splitted into:
    ['AB-Betrieb', 'm', 'Đ_TỬ', 'dass AB', 'mode', 'chê'độ hạng', 'AB']
    """
    newword_texts = []
    newwordcase_capitals = []
    newword_format_bolds = []
    newword_format_italics = []
    for k in range(len(word_texts)):
        phrase = word_texts[k]
        words, wordcase_capital = split_capital(phrase)
        newword_texts.extend(words)
        newwordcase_capitals.extend(wordcase_capital)
        for n in range(len(words)):
            newword_format_bolds.append(word_format_bolds[k])
            newword_format_italics.append(word_format_italics[k])
    newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = \
    merge_with_comment_phrase(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
    newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = \
    merge_similar_series(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
    return newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics

# encoding=utf8 by default
# See: https://stackoverflow.com/questions/21129020/how-to-fix-unicodedecodeerror-ascii-codec-cant-decode-byte

def WriteTableCsv(exportFile, table):
    with open(exportFile, 'wt') as csvfile:
        textWriter = csv.writer(csvfile, delimiter=str('\t'))
        for i in range(len(table)):
            textWriter.writerow(table[i])

def analyze_line(word_texts, wordcase_capitals, word_format_bolds, word_format_italics):
    """Algorithm to analyze a line:
    A line needs at least 4 words. If not,  neglect the line.
    First word is considered the main word, in German.
    Second word is considered the type of the German word.
    Third word must be CAPITAL, indicating the FIELD. If not, neglect the line.
    After a CAPITAL word: lowercase non-italic words are English, lowercase italic words are Vietnamese.
    Split items for different FIELDS.
    Split items for new English meaning (the precedent word is Vietnamese).
    """
    de_word = []
    en_word = []
    vi_word = []
    type_word = []
    field_word = []
    if len(word_texts) < 4:
        out_message = 'Line has less than 4 elements, skipped.'
        if len(word_texts) >= 1:
            out_message += ' This line starts with: ' + word_texts[0]
        return de_word, type_word, field_word, en_word, vi_word, out_message
    if wordcase_capitals[2] != True:
        out_message = '3rd word is not CAPITAL, syntax not compatible, skipped.'
        out_message += ' This line starts with: ' + word_texts[0]
        return de_word, type_word, field_word, en_word, vi_word, out_message
    item = 0
    de_word.append(word_texts[0])
    type_word.append(word_texts[1])
    field_word.append(word_texts[2])
    en_word.append('')
    vi_word.append('')
    for k in range(3, len(word_texts)):
        if wordcase_capitals[k] == True:
            item += 1
            de_word.append(word_texts[0])
            type_word.append(word_texts[1])
            field_word.append(word_texts[k])
            en_word.append('')
            vi_word.append('')
        else:
            if word_format_italics[k] != True:
                if word_format_italics[k-1] == True and wordcase_capitals[k-1] != True:
                    item += 1
                    de_word.append(word_texts[0])
                    type_word.append(word_texts[1])
                    field_word.append(field_word[item-1])
                    en_word.append('')
                    en_word[item] += ' '+word_texts[k]
                    vi_word.append('')
                else:
                    en_word[item] += ' '+word_texts[k]
            else:
                vi_word[item] += ' '+word_texts[k]
    de_word = [word.strip().strip(',') for word in de_word]
    en_word = [word.strip().strip(',') for word in en_word]
    vi_word = [word.strip().strip(',') for word in vi_word]
    type_word = [word.strip().strip(',') for word in type_word]
    field_word = [word.strip().strip(',') for word in field_word]
    out_message = 'Line has been analyzed successfully to ' + str(item+1) + ' item(s).'
    for k in range(len(de_word)):
        if en_word[k] == '' or vi_word[k] == '':
            out_message += ' Attention: this entry does misses English or Vietnamese word(s): ' + de_word[k]
    return de_word, type_word, field_word, en_word, vi_word, out_message

## Tests:
##line = doc.paragraphs[2]
##line = doc.paragraphs[16]
##line = doc.paragraphs[19]
##line = doc.paragraphs[25]
#
#line = doc.paragraphs[7]
##for phrase in line.runs:
##   print(phrase.text)
##   print(line.style.name)
##   print(line.style.font.bold)
##   print(line.style.font.italic)
##   print(phrase.style.name)
##   print(phrase.style.font.bold)
##   print(phrase.style.font.italic)
#
#paragraph_style_bold = line.style.font.bold
#paragraph_style_italic = line.style.font.italic
#word_texts = [part.text.strip() for part in line.runs if part.text.strip() != '']
#word_style_names = [part.style.name for part in line.runs if part.text.strip() != '']
#word_style_bolds = [part.style.font.bold for part in line.runs if part.text.strip() != '']
#word_style_italics = [part.style.font.italic for part in line.runs if part.text.strip() != '']
#word_format_bolds, word_format_italics = read_format(word_style_bolds, word_style_italics, paragraph_style_bold, paragraph_style_italic)
#
#for k in range(len(word_style_names)):
#    print(word_texts[k])
#    print(paragraph_style_bold)
#    print(paragraph_style_italic)
#    print(word_style_names[k])
#    print(word_format_bolds[k])
#    print(word_format_italics[k])
#
#newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = re_parse(word_texts, word_format_bolds, word_format_italics)
#
#for k in range(len(newword_texts)):
#    print(newword_texts[k])
#    print(newwordcase_capitals[k])
#    print(newword_format_bolds[k])
#    print(newword_format_italics[k])
#
#de_word, type_word, field_word, en_word, vi_word, out_message = analyze_line(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
#
#print(de_word)
#print(type_word)
#print(field_word)
#print(en_word)
#print(vi_word)

# Main operation:
def readocr(inputFile, exportFile = 'result.csv', logFile = 'log.txt'):

    de_words = []
    en_words = []
    vi_words = []
    type_words = []
    field_words = []
    out_messages = []
    
    doc = docx.Document(inputFile)

    for line in doc.paragraphs:
        paragraph_style_bold = line.style.font.bold
        paragraph_style_italic = line.style.font.italic
        word_texts = [part.text.strip() for part in line.runs if part.text.strip() != '']
        word_style_bolds = [part.style.font.bold for part in line.runs if part.text.strip() != '']
        word_style_italics = [part.style.font.italic for part in line.runs if part.text.strip() != '']
        word_format_bolds, word_format_italics = read_format(word_style_bolds, word_style_italics, paragraph_style_bold, paragraph_style_italic)
        #print(word_texts)
        newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = re_parse(word_texts, word_format_bolds, word_format_italics)
        #print(newword_texts)    
        de_word, type_word, field_word, en_word, vi_word, out_message = analyze_line(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
        #print(de_word)
        #print(en_word)
        #print(vi_word)
        de_words.extend(de_word)
        en_words.extend(en_word)
        vi_words.extend(vi_word)
        type_words.extend(type_word)
        field_words.extend(field_word)
        out_messages.append(out_message)
    
    # Note the meaning of line.runs[0].style.font.bold: (and line.runs[0].style.font.italic)
        # line.runs[0].style.font.bold is "toggle" for line.style.font.bold, i.e.
        # if line.style.font.bold == True, and line.runs[0].style.font.bold == True
        # then the real style is False. Probably the XOR operator of the two booleans works. 
    
    # Later: output directly the result from each line to the CSV file
    table = []
    for k in range(len(de_words)):
        table.append([de_words[k], type_words[k], field_words[k], en_words[k], vi_words[k]])

    WriteTableCsv(exportFile, table)
    with open(logFile, 'wt') as txtfile:
        for item in out_messages:
            txtfile.write("%s\n" % item)


