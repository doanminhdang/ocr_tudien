#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Parse a DOCX file to get dictionary words and output glossary table

Created on Mon May 14 2018

@author: dang

Written for Python 2.6
"""

from __future__ import unicode_literals
import docx
import sys
reload(sys)
sys.setdefaultencoding('utf8')
import csv
import re

def check_series(text_list, set_list):
    """Check whether all items in each text_list[k] are in set_list
    """
    in_list = []
    for word in text_list:
        all_words = re.sub('\(.*?\)', ',', word).split(',')
        all_words = list(filter(None, all_words))
        component_in_list = [component.strip(' ') in set_list for component in all_words]
        this_word_in_list = all(component_in_list)
        in_list.append(this_word_in_list)
    return in_list

def check_field(text_list):
    with open('../ocr_tudien/word_fields.txt', 'rt') as listfile:
        set_list = listfile.read().split('\n')
        if set_list[-1] == '':
            del(set_list[-1])
    in_list = check_series(text_list, set_list)
    return in_list

def check_type(text_list):
    with open('../ocr_tudien/word_types.txt', 'rt') as listfile:
        set_list = listfile.read().split('\n')
        if set_list[-1] == '':
            del(set_list[-1])
    in_list = check_series(text_list, set_list)
    return in_list

def read_format(paragraph_style_bold, character_style_bolds, character_font_bolds, paragraph_style_italic, character_style_italics, character_font_italics):
    """Detect the character properties (bold, italic) based on info of the hierarchy:
       paragraph_style_bold (single value, from line.style.font.bold): on top level of the hierarchy, can be "None", "True", or "False"
       character_style_bolds (array, from docx line.runs[0].style.font.bold): second level, it inherit upper level if "None", override if "True" or "False"
       character_font_bolds (array, from docx line.runs[0].font.bold): third level, directly applied to character if "True" or "False", inherit if "None"
    """
    # Using style name doesn't work with different languages, example:
    # English Word file: Bold, Not Italic, German: Fett, Nicht kursiv
    if paragraph_style_bold == None:
        word_format_bolds = [False for item in character_font_bolds]
    else:
        word_format_bolds = [paragraph_style_bold for item in character_font_bolds]
    for k in range(len(character_style_bolds)):

        if character_style_bolds[k] != None:
            word_format_bolds[k] = character_style_bolds[k]
        if character_font_bolds[k] != None:
            word_format_bolds[k] = character_font_bolds[k]
    
    if paragraph_style_italic == None:
        word_format_italics = [False for item in character_font_italics]
    else:
        word_format_italics = [paragraph_style_italic for item in character_font_italics]
    for k in range(len(character_style_italics)):
        if character_style_italics[k] != None:
            word_format_italics[k] = character_style_italics[k]
        if character_font_italics[k] != None:
            word_format_italics[k] = character_font_italics[k]
    
    return word_format_bolds, word_format_italics

def remove_empty(list_text, *list_properties):
    for k in range(len(list_text)-1, 0, -1):
        if list_text[k] == '':
            del(list_text[k])
            for prop in list_properties:
                del(prop[k])
    
def split_capital(phrase):
    """Split capital and small words.
    Capital words are only allowed at the first half of the phrase.
    If a capital word is after a small word, it is NOT marked capital.
    """
    format_capital = []
    #words = [phrase]
    #format_capital.append(words[0].isupper())
    
    #words = phrase.split()
    re_split_result = re.split('(\W)', phrase)
    words = list(filter(None, re_split_result))
    
    for k_item in range(len(words)-1, 0, -1):
            while len(words[k_item])>0 and words[k_item][0] == ',':
                words[k_item-1] += ','
                words[k_item] = words[k_item][1:]
    for k_item in range(len(words)-1, 0, -1):
        while len(words[k_item])>0 and words[k_item][0] == ' ':
            words[k_item-1] += ' '
            words[k_item] = words[k_item][1:]
        
    remove_empty(words)
        
    for k_item in range(len(words)-1, 0, -1):
        if words[k_item][0] != ' ' and words[k_item-1][-1] != ' ':
            words[k_item-1] = ''.join([words[k_item-1], words[k_item]])
            del(words[k_item])
    
    flag_prev_word_capital = True
    for k in range(len(words)):
        if words[k].isupper() and flag_prev_word_capital:
            if k>0 and (words[k-1].rstrip(' ')[-1] != ',' and words[k-1].rstrip(' ')[-1] != ')'):
                format_capital.append(False)
                flag_prev_word_capital = False
            else:
                format_capital.append(True)
            
        else:
            format_capital.append(False)
            flag_prev_word_capital = False
    return words, format_capital

def merge_similar_series(words, *list_properties):
    for k in range(len(words)-1,0,-1):
        same_flag = True
        for type_format in list_properties:
            same_flag = same_flag and (type_format[k] == type_format[k-1])
        if same_flag:
            words[k-1] = ''.join([words[k-1], words[k]])
            del words[k]
            for type_format in list_properties:
                del type_format[k]

def merge_with_comment_phrase(words, *list_properties):
    """If the starting word in phrase is in (), it should be a comment for the previous one,
    hence be added to the previous one.
    """
    for k in range(len(words)-1,0,-1):
        if re.match('^(\(.*?\))', words[k].strip()):
            words[k-1] = ''.join([words[k-1], re.search('( *\(.*?\))', words[k].strip()).groups()[0]])
            words[k] = re.sub('(\(.*?\))', '', words[k])
            if words[k].strip(', ') == '':
                words[k-1] = ''.join([words[k-1], words[k]])
                del words[k]
                for type_format in list_properties:
                    del type_format[k]

def re_parse(word_texts, word_format_bolds, word_format_italics):
    """Fix the leftover mistake in the data parsed, some example items:
    THAN gate road
    TOÁN, S_CHÊ figure 
    hình 
    (bản vẽ) - is a separate item, in fact it is a comment for previous item
    AB-Betrieb m Đ_TỬ dass AB mode chê'độ hạng AB, splitted into:
    ['AB-Betrieb', 'm', 'Đ_TỬ', 'dass AB', 'mode', 'chê'độ hạng', 'AB']
    """
    merge_with_comment_phrase(word_texts, word_format_bolds, word_format_italics)
    merge_similar_series(word_texts, word_format_bolds, word_format_italics)
    newword_texts = []
    newwordcase_capitals = []
    newword_format_bolds = []
    newword_format_italics = []
    # first word is assumed German, do not need to split capital
    if len(word_texts)>0:
        newword_texts.append(word_texts[0])
        newwordcase_capitals.append(False)
        newword_format_bolds.append(word_format_bolds[0])
        newword_format_italics.append(word_format_italics[0])
    for k in range(1, len(word_texts)):
        if word_format_italics[k]:
            words = [word_texts[k]]
            wordcase_capital = [False]
        else:
            words, wordcase_capital = split_capital(word_texts[k])
        newword_texts.extend(words)
        newwordcase_capitals.extend(wordcase_capital)
        for n in range(len(words)):
            newword_format_bolds.append(word_format_bolds[k])
            newword_format_italics.append(word_format_italics[k])
    
    merge_with_comment_phrase(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
    merge_similar_series(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
    return newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics

# encoding=utf8 by default
# See: https://stackoverflow.com/questions/21129020/how-to-fix-unicodedecodeerror-ascii-codec-cant-decode-byte

def WriteTableCsv(exportFile, table):
    with open(exportFile, 'wt') as csvfile:
        textWriter = csv.writer(csvfile, delimiter=str('\t'))
        for i in range(len(table)):
            textWriter.writerow(table[i])

def analyze_line(word_texts, wordcase_capitals, word_format_bolds, word_format_italics):
    """Algorithm to analyze a line:
    A line needs at least 4 words. If not,  neglect the line.
    First word is considered the main word, in German.
    Second word is considered the type of the German word.
    Third word must be CAPITAL, indicating the FIELD. If not, neglect the line.
    After a CAPITAL word: lowercase non-italic words are English, lowercase italic words are Vietnamese.
    Split items for different FIELDS.
    Split items for new English meaning (the precedent word is Vietnamese).
    
    Some special cases: 
    ABC v_tắt (automatische Helligkeitsregelung) TV ABC (automatic brightness control) (sự) điều chỉnh độ chói tự động
    """
    de_word = []
    en_word = []
    vi_word = []
    type_word = []
    field_word = []
    out_message = ''
    #out_message = '-'.join(word_texts)
    #out_message += '\nCapital: '
    #out_message += '-'.join([str(item) for item in wordcase_capitals])
    #out_message += '\nBold: '
    #out_message += '-'.join([str(item) for item in word_format_bolds])
    #out_message += '\nItalic: '
    #out_message += '-'.join([str(item) for item in word_format_italics])
    #out_message += '\n'
    
    if ''.join(word_texts).strip() == '':
        out_message += 'Blank line, skipped.'
        return de_word, type_word, field_word, en_word, vi_word, out_message
    if len(word_texts) < 4:
        out_message += 'ERROR: Line has less than 4 elements, skipped.'
        if len(word_texts) >= 1:
            out_message += ' This line starts with: ' + word_texts[0]
        return de_word, type_word, field_word, en_word, vi_word, out_message
    if wordcase_capitals[2] != True:
        out_message += 'ERROR: 3rd word is not CAPITAL, syntax not compatible, skipped.'
        out_message += ' This line starts with: ' + word_texts[0]
        return de_word, type_word, field_word, en_word, vi_word, out_message
    item = 0
    de_word.append(word_texts[0])
    type_word.append(word_texts[1])
    field_word.append(word_texts[2])
    en_word.append('')
    vi_word.append('')
    for k in range(3, len(word_texts)):
        if wordcase_capitals[k] == True:
            item += 1
            de_word.append(word_texts[0])
            type_word.append(word_texts[1])
            field_word.append(word_texts[k])
            en_word.append('')
            vi_word.append('')
        else:
            if word_format_italics[k] != True:
                if word_format_italics[k-1] == True and wordcase_capitals[k-1] != True:
                    item += 1
                    de_word.append(word_texts[0])
                    type_word.append(word_texts[1])
                    field_word.append(field_word[item-1])
                    en_word.append('')
                    en_word[item] += ' '+word_texts[k]
                    vi_word.append('')
                else:
                    en_word[item] += ' '+word_texts[k]
            else:
                vi_word[item] += ' '+word_texts[k]
    de_word = [word.strip().strip(',') for word in de_word]
    en_word = [word.strip().strip(',') for word in en_word]
    vi_word = [word.strip().strip(',') for word in vi_word]
    type_word = [word.strip().strip(',') for word in type_word]
    field_word = [word.strip().strip(',') for word in field_word]
    out_message += 'Line has been analyzed successfully to ' + str(item+1) + ' item(s).'
    return de_word, type_word, field_word, en_word, vi_word, out_message

## Tests:
##import docx
##inputFile = 'abc.docx'
##doc = docx.Document(inputFile)
##line = doc.paragraphs[2]
##line = doc.paragraphs[16]
##line = doc.paragraphs[19]
##line = doc.paragraphs[25]
#
#line = doc.paragraphs[7]
#print line.text
#for phrase in line.runs:
  #print(phrase.text)
  #print(line.style.name)
  #print(line.style.font.bold)
  #print(line.style.font.italic)
  #print(phrase.style.name)
  #print(phrase.style.font.name)
  #print(phrase.style.font.bold)
  #print(phrase.style.font.italic)
#
#paragraph_style_bold = line.style.font.bold
#paragraph_style_italic = line.style.font.italic
#word_texts = [part.text.strip() for part in line.runs if part.text.strip() != '']
#word_style_names = [part.style.name for part in line.runs if part.text.strip() != '']
#word_style_bolds = [part.style.font.bold for part in line.runs if part.text.strip() != '']
#word_style_italics = [part.style.font.italic for part in line.runs if part.text.strip() != '']
#word_format_bolds, word_format_italics = read_format(word_style_bolds, word_style_italics, paragraph_style_bold, paragraph_style_italic)
#
#for k in range(len(word_style_names)):
#    print(word_texts[k])
#    print(paragraph_style_bold)
#    print(paragraph_style_italic)
#    print(word_style_names[k])
#    print(word_format_bolds[k])
#    print(word_format_italics[k])
#
#newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = re_parse(word_texts, word_format_bolds, word_format_italics)
#
#for k in range(len(newword_texts)):
#    print(newword_texts[k])
#    print(newwordcase_capitals[k])
#    print(newword_format_bolds[k])
#    print(newword_format_italics[k])
#
#de_word, type_word, field_word, en_word, vi_word, out_message = analyze_line(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
#
#print(de_word)
#print(type_word)
#print(field_word)
#print(en_word)
#print(vi_word)

def analyze_line_tudien_sinhhoc(word_texts, wordcase_capitals, word_format_bolds, word_format_italics):
    """Algorithm to analyze a line in Tu dien Sinh hoc:
    A line needs at least 2 words. If not,  neglect the line.
    First phrase non-italic words is the English words.
    Second phrase, italic, is the Vietnamse words.
    Split items for new Vietnamese meaning (started by 1. or 2. ...).
    
    Example:
    aspect 1. quang cảnh 2. sắc thái
    TODO: Special case, not yet parsed correctly, still need to fix!
    achondroplastic (thuộc) (chứng) loạn sản sụn
    """
    en_word = []
    vi_word = []
    out_message = ''
    
    if ''.join(word_texts).strip() == '':
        out_message += 'Blank line, skipped.'
        return en_word, vi_word, out_message
    if len(word_texts) < 2:
        out_message += 'ERROR: Line has less than 2 elements, skipped.'
        if len(word_texts) >= 1:
            out_message += ' This line starts with: ' + word_texts[0]
        return en_word, vi_word, out_message
    item = 0
    en_word.append('')
    vi_word.append('')
    vi_word_flag = False
    for k in range(len(word_texts)):
        if word_format_italics[k] != True:
            if vi_word_flag != False:
                item += 1
                en_word.append(word_texts[k])
                vi_word.append('')
                vi_word_flag = False
            else:
                en_word[item] += ' '+word_texts[k]
        else:
            vi_word[item] += ' '+word_texts[k]
            vi_word_flag = True
            # Parse multiple Vietnamese words, example: t='1. miệng giác 2. lỗ chân 3. ổ khớp 4. múi nhau'
            if vi_word[item].find('1.') != -1:
                temp_vi_words = vi_word[item]
                number=1
                old_number_pos=temp_vi_words.find(str(number)+'.')
                number+=1
                number_pos=temp_vi_words.find(str(number)+'.')
                while number_pos != -1:
                    # when the next number (== Vietnamese meaning) is available: split this meaning
                    vi_word[item] = temp_vi_words[old_number_pos+2:number_pos]
                    item += 1
                    en_word.append(en_word[item-1])
                    vi_word.append('')
                    old_number_pos=number_pos
                    number+=1
                    number_pos=temp_vi_words.find(str(number)+'.')
                vi_word[item] = temp_vi_words[old_number_pos+2:]  # the last part
            
    en_word = [word.strip().strip(',') for word in en_word]
    vi_word = [word.strip().strip(',') for word in vi_word]
    out_message += 'Line has been analyzed successfully to ' + str(item+1) + ' item(s).'
    return en_word, vi_word, out_message

def analyze_line_physics_NDH(word_texts, wordcase_capitals, word_format_bolds, word_format_italics):
    """Algorithm to analyze a line in Tu dien Sinh hoc:
    A line needs at least 2 words. If not,  neglect the line.
    First phrase non-italic words is the English words.
    Second phrase, italic, is the Vietnamse words.
    Split items for new Vietnamese meaning (started by 1. or 2. ...).
    
    Example:
    wavelength /ˈweɪvˌlɛŋkθ/ (n): bước sóng
    """
    en_word = []
    vi_word = []
    pronunciation_word = []
    type_word = []
    out_message = ''
    
    if ''.join(word_texts).strip() == '':
        out_message += 'Blank line, skipped.'
        return en_word, vi_word, pronunciation_word, type_word, out_message
    if len(word_texts) < 2:
        out_message += 'ERROR: Line has less than 2 elements, skipped.'
        if len(word_texts) >= 1:
            out_message += ' This line starts with: ' + word_texts[0]
        return en_word, vi_word, pronunciation_word, type_word, out_message
    item = 0
    en_word.append('')
    vi_word.append('')
    pronunciation_word.append('')
    type_word.append('')
    
    for k in range(len(word_texts)):
        temp_rest_words = ''
        if word_format_bolds[k] == True:
            en_word[item] += ' '+word_texts[k]
        else:
            temp_rest_words += word_texts[k]
        # Parse the rest of the words: 'pronunciation_word (type): Vietnamese'
        vi_word[item] += temp_rest_words[temp_rest_words.find(':')+2:]
        if re.search(r'/.*?/', temp_rest_words):
            pronunciation_word[item] += re.search(r'/.*?/', temp_rest_words).group(0)
        if re.search(r'\(.*?\).*:',temp_rest_words):
            type_word[item] += re.search(r'\(.*?\).*:',temp_rest_words).group(0)[:-1]  # do not keep the final :
            
    en_word = [word.strip().strip(',') for word in en_word]
    vi_word = [word.strip().strip(',') for word in vi_word]
    pronunciation_word = [word.strip('/') for word in pronunciation_word]
    type_word = [word.strip('()').strip() for word in type_word]
    
    out_message += 'Line has been analyzed successfully to ' + str(item+1) + ' item(s).'
    return en_word, vi_word, pronunciation_word, type_word, out_message

def readocr(inputFile, exportFile = 'result.csv', logFile = 'log.txt'):

    de_words = []
    en_words = []
    vi_words = []
    type_words = []
    field_words = []
    out_messages = []
    
    number_items_in_line = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
    
    number_blank_lines = 0
    number_type_issues = 0
    number_field_issues = 0
    number_en_issues = 0
    number_vi_issues = 0
    number_total_item_issues = 0
    
    doc = docx.Document(inputFile)

    for k in range(len(doc.paragraphs)):
        line = doc.paragraphs[k]
        # Make it compatible to Wordpad, this software removes all the styles
        paragraph_style_bold = None
        paragraph_style_italic = None
        if line.style != None:
            paragraph_style_bold = line.style.font.bold
            paragraph_style_italic = line.style.font.italic
        word_texts = [part.text for part in line.runs]
        character_style_bolds = [None for part in line.runs]
        character_style_italics = [None for part in line.runs]
        for k_item in range(len(line.runs)):
            part = line.runs[k_item]
            if part.style != None:
                character_style_bolds[k_item] = part.style.font.bold
                character_style_italics[k_item] = part.style.font.italic
        character_font_bolds = [part.font.bold for part in line.runs]
        character_font_italics = [part.font.italic for part in line.runs]
        
        remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == '&':
                word_texts[k_item-1] += '&'
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ',':
                word_texts[k_item-1] += ','
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ' ':
                word_texts[k_item-1] += ' '
                word_texts[k_item] = word_texts[k_item][1:]
        
        remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        for k_item in range(len(word_texts)-1, 0, -1):
            if word_texts[k_item][0] != ' ' and word_texts[k_item-1][-1] != ' ':
                word_texts[k_item-1] = ''.join([word_texts[k_item-1], word_texts[k_item]])
                del(word_texts[k_item])
                del(character_style_bolds[k_item])
                del(character_style_italics[k_item])
                del(character_font_bolds[k_item])
                del(character_font_italics[k_item])
            
        word_format_bolds, word_format_italics = read_format(paragraph_style_bold, character_style_bolds, character_font_bolds, paragraph_style_italic, character_style_italics, character_font_italics)
        #print(word_texts)
        newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = re_parse(word_texts, word_format_bolds, word_format_italics)
        #print(newword_texts)    
        de_word, type_word, field_word, en_word, vi_word, out_message = analyze_line(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
        #print(de_word)
        #print(en_word)
        #print(vi_word)
        de_words.extend(de_word)
        en_words.extend(en_word)
        vi_words.extend(vi_word)
        type_words.extend(type_word)
        field_words.extend(field_word)
        
        # Accumulative statistics
        if line.text.strip() == '':
            number_blank_lines += 1
        number_items_in_line[len(de_word)%11] += 1
        
        is_good_type = check_type(type_word)
        is_good_field = check_field(field_word)
        for k_entry in range(len(de_word)):
            if is_good_type[k_entry] != True or is_good_field[k_entry] != True or en_word[k_entry] == '' or vi_word[k_entry] == '':
                number_total_item_issues += 1
                out_message += '\nAttention - suspected problem with this item: ' + de_word[k_entry] + ' - ' + type_word[k_entry] + ' - ' + field_word[k_entry] + ' - ' + en_word[k_entry] + ' - ' + vi_word[k_entry] + '.'
            if is_good_type[k_entry] != True:
                number_type_issues += 1
                out_message += '\n Check the type word.'
            if is_good_field[k_entry] != True:
                number_field_issues += 1
                out_message += '\n   Check the FIELD word.'
            if en_word[k_entry] == '':
                number_en_issues += 1
                out_message += '\n     English word is missing.'
            if vi_word[k_entry] == '':
                number_vi_issues += 1
                out_message += '\n       Vietnamese word is missing.'
        
        out_messages.append('Line ' + str(k+1) + ': ' + out_message)
        
        # Note the meaning of line.style.font.bold, line.runs[0].style.font.bold, and line.runs[0].font.bold:
        # line.style.font.bold: on top level of the hierarchy
        # line.runs[0].style.font.bold: second level, it inherit upper level if "None", override if "True" or "False"
        # line.runs[0].font.bold: third level, directly applied to character if "True" or "False", inherit if "None"
        #
        # http://python-docx.readthedocs.io/en/latest/user/styles-using.html
        # Many font properties are tri-state, meaning they can take the values True, False, and None.
        # True means the property is “on”, False means it is “off”.
        # Conceptually, the None value means “inherit”. 
        # Because a style exists in an inheritance hierarchy, it is important 
        # to have the ability to specify a property at the right place in the hierarchy,
        # generally as far up the hierarchy as possible.
        # For example, if all headings should be in the Arial typeface,
        # it makes more sense to set that property on the Heading 1 style
        # and have Heading 2 inherit from Heading 1.
    
    # Later: output directly the result from each line to the CSV file
    table = []
    for k in range(len(de_words)):
        table.append([de_words[k], type_words[k], field_words[k], en_words[k], vi_words[k]])

    WriteTableCsv(exportFile, table)
    with open(logFile, 'wt') as txtfile:
        txtfile.write("Total number of lines processed: %d\n" % len(doc.paragraphs))
        txtfile.write("  with number of blank lines: %d\n" % number_blank_lines)
        txtfile.write("  \n*Number of lines failed to process (0 item): %d\n\n" % number_items_in_line[0])
        txtfile.write("  Number of lines with 1 item: %d\n" % number_items_in_line[1])
        txtfile.write("  Number of lines with 2 items: %d\n" % number_items_in_line[2])
        txtfile.write("  Number of lines with 3 items: %d\n" % number_items_in_line[3])
        txtfile.write("  Number of lines with 4 items: %d\n" % number_items_in_line[4])
        txtfile.write("  Number of lines with 5 items: %d\n" % number_items_in_line[5])
        txtfile.write("  Number of lines with 6 items: %d\n" % number_items_in_line[6])
        txtfile.write("  Number of lines with 7 items: %d\n" % number_items_in_line[7])
        txtfile.write("  Number of lines with 8 items: %d\n" % number_items_in_line[8])
        txtfile.write("  Number of lines with 9 items: %d\n" % number_items_in_line[9])
        txtfile.write("  Number of lines with 10 items: %d\n" % number_items_in_line[10])
                
        txtfile.write("\nTotal items obtained: %d\n" % len(de_words))
        txtfile.write("\n*Number of items with issues: %d\n" % number_total_item_issues)
        txtfile.write("  in which, there are:\n")
        txtfile.write("    + %d items with awkward type\n" % number_type_issues)
        txtfile.write("    + %d items with awkward field\n" % number_field_issues)
        txtfile.write("    + %d items without English word\n" % number_en_issues)
        txtfile.write("    + %d items without Vietnamese word\n" % number_vi_issues)
        txtfile.write("\n-----\n\n")
        
        for item in out_messages:
            txtfile.write("%s\n" % item)
    
    #excelFile = exportFile[:-4] + '.xls'
    #Csv_Excel.csv_to_xls(exportFile, excelFile)

def readocr_tudien_sinhhoc(inputFile, exportFile = 'result.csv', logFile = 'log.txt'):
    
    with open('tdshav_letters_incorrect.txt', 'rt') as search_file:
        search_text = search_file.readlines()
    search_text = [item.strip('\n') for item in search_text]
    
    with open('tdshav_letters_correct.txt', 'rt') as replace_file:
        replace_text = replace_file.readlines()
    replace_text = [item.strip('\n') for item in replace_text]
    

    en_words = []
    vi_words = []
    out_messages = []
    
    number_items_in_line = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
    
    number_blank_lines = 0
    number_en_issues = 0
    number_vi_issues = 0
    number_total_item_issues = 0
    
    doc = docx.Document(inputFile)

    for k in range(len(doc.paragraphs)):
        line = doc.paragraphs[k]
        # Make it compatible to Wordpad, this software removes all the styles
        paragraph_style_bold = None
        paragraph_style_italic = None
        if line.style != None:
            paragraph_style_bold = line.style.font.bold
            paragraph_style_italic = line.style.font.italic
        word_texts = [part.text for part in line.runs]
        character_style_bolds = [None for part in line.runs]
        character_style_italics = [None for part in line.runs]
        for k_item in range(len(line.runs)):
            part = line.runs[k_item]
            if part.style != None:
                character_style_bolds[k_item] = part.style.font.bold
                character_style_italics[k_item] = part.style.font.italic
        character_font_bolds = [part.font.bold for part in line.runs]
        character_font_italics = [part.font.italic for part in line.runs]
        
        remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == '&':
                word_texts[k_item-1] += '&'
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ',':
                word_texts[k_item-1] += ','
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ' ':
                word_texts[k_item-1] += ' '
                word_texts[k_item] = word_texts[k_item][1:]
        
        remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        for k_item in range(len(word_texts)-1, 0, -1):
            if word_texts[k_item][0] != ' ' and word_texts[k_item-1][-1] != ' ':
                word_texts[k_item-1] = ''.join([word_texts[k_item-1], word_texts[k_item]])
                del(word_texts[k_item])
                del(character_style_bolds[k_item])
                del(character_style_italics[k_item])
                del(character_font_bolds[k_item])
                del(character_font_italics[k_item])
        
        # replace the series of letters for Vietnamese
        for k_letter in range(len(search_text)):
            word_texts = [text_item.replace(search_text[k_letter], replace_text[k_letter]) for text_item in word_texts]
        
        word_format_bolds, word_format_italics = read_format(paragraph_style_bold, character_style_bolds, character_font_bolds, paragraph_style_italic, character_style_italics, character_font_italics)
        #print(word_texts)
        newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = re_parse(word_texts, word_format_bolds, word_format_italics)
        #print(newword_texts)
        en_word, vi_word, out_message = analyze_line_tudien_sinhhoc(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
        #print(en_word)
        #print(vi_word)
        en_words.extend(en_word)
        vi_words.extend(vi_word)
        
        # Accumulative statistics
        if line.text.strip() == '':
            number_blank_lines += 1
        number_items_in_line[len(en_word)%11] += 1

        for k_entry in range(len(en_word)):
            if en_word[k_entry] == '' or vi_word[k_entry] == '':
                number_total_item_issues += 1
                out_message += '\nAttention - suspected problem with this item: ' + en_word[k_entry] + ' - ' + vi_word[k_entry] + '.'
            if en_word[k_entry] == '':
                number_en_issues += 1
                out_message += '\n     English word is missing.'
            if vi_word[k_entry] == '':
                number_vi_issues += 1
                out_message += '\n       Vietnamese word is missing.'
        
        out_messages.append('Line ' + str(k+1) + ': ' + out_message)
        
    # Later: output directly the result from each line to the CSV file
    table = []
    for k in range(len(en_words)):
        table.append([en_words[k], vi_words[k]])

    WriteTableCsv(exportFile, table)
    with open(logFile, 'wt') as txtfile:
        txtfile.write("Total number of lines processed: %d\n" % len(doc.paragraphs))
        txtfile.write("  with number of blank lines: %d\n" % number_blank_lines)
        txtfile.write("  \n*Number of lines failed to process (0 item): %d\n\n" % number_items_in_line[0])
        txtfile.write("  Number of lines with 1 item: %d\n" % number_items_in_line[1])
        txtfile.write("  Number of lines with 2 items: %d\n" % number_items_in_line[2])
        txtfile.write("  Number of lines with 3 items: %d\n" % number_items_in_line[3])
        txtfile.write("  Number of lines with 4 items: %d\n" % number_items_in_line[4])
        txtfile.write("  Number of lines with 5 items: %d\n" % number_items_in_line[5])
        txtfile.write("  Number of lines with 6 items: %d\n" % number_items_in_line[6])
        txtfile.write("  Number of lines with 7 items: %d\n" % number_items_in_line[7])
        txtfile.write("  Number of lines with 8 items: %d\n" % number_items_in_line[8])
        txtfile.write("  Number of lines with 9 items: %d\n" % number_items_in_line[9])
        txtfile.write("  Number of lines with 10 items: %d\n" % number_items_in_line[10])
                
        txtfile.write("\nTotal items obtained: %d\n" % len(en_words))
        txtfile.write("\n*Number of items with issues: %d\n" % number_total_item_issues)
        txtfile.write("  in which, there are:\n")
        txtfile.write("    + %d items without English word\n" % number_en_issues)
        txtfile.write("    + %d items without Vietnamese word\n" % number_vi_issues)
        txtfile.write("\n-----\n\n")
        
        for item in out_messages:
            txtfile.write("%s\n" % item)

def readocr_physics_NDH(inputFile, exportFile = 'result.csv', logFile = 'log.txt'):
    """Parse the mini physics dictionary of Nguyen Dong Hai
    """

    en_words = []
    vi_words = []
    pronunciation_words = []
    type_words = []
    out_messages = []
    
    number_items_in_line = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
    
    number_blank_lines = 0
    number_en_issues = 0
    number_vi_issues = 0
    number_pronunciation_issues = 0
    number_type_issues = 0
    number_total_item_issues = 0
    
    doc = docx.Document(inputFile)

    for k in range(len(doc.paragraphs)):
        line = doc.paragraphs[k]
        # Make it compatible to Wordpad, this software removes all the styles
        paragraph_style_bold = None
        paragraph_style_italic = None
        if line.style != None:
            paragraph_style_bold = line.style.font.bold
            paragraph_style_italic = line.style.font.italic
        word_texts = [part.text for part in line.runs]
        character_style_bolds = [None for part in line.runs]
        character_style_italics = [None for part in line.runs]
        for k_item in range(len(line.runs)):
            part = line.runs[k_item]
            if part.style != None:
                character_style_bolds[k_item] = part.style.font.bold
                character_style_italics[k_item] = part.style.font.italic
        character_font_bolds = [part.font.bold for part in line.runs]
        character_font_italics = [part.font.italic for part in line.runs]
        
        remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == '&':
                word_texts[k_item-1] += '&'
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ',':
                word_texts[k_item-1] += ','
                word_texts[k_item] = word_texts[k_item][1:]
        for k_item in range(len(word_texts)-1, 0, -1):
            while len(word_texts[k_item])>0 and word_texts[k_item][0] == ' ':
                word_texts[k_item-1] += ' '
                word_texts[k_item] = word_texts[k_item][1:]
        
        remove_empty(word_texts, character_style_bolds, character_style_italics, character_font_bolds, character_font_italics)
        
        for k_item in range(len(word_texts)-1, 0, -1):
            #print k_item
            #print 'item'+word_texts[k_item]+'item'
            #print 'item-1'+word_texts[k_item-1]+'item-1'
            if word_texts[k_item][0] != ' ':
                if word_texts[k_item-1] == '' or word_texts[k_item-1][-1] != ' ':
                    word_texts[k_item-1] = ''.join([word_texts[k_item-1], word_texts[k_item]])
                    del(word_texts[k_item])
                    del(character_style_bolds[k_item])
                    del(character_style_italics[k_item])
                    del(character_font_bolds[k_item])
                    del(character_font_italics[k_item])
            
        word_format_bolds, word_format_italics = read_format(paragraph_style_bold, character_style_bolds, character_font_bolds, paragraph_style_italic, character_style_italics, character_font_italics)
        #print(word_texts)
        newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics = re_parse(word_texts, word_format_bolds, word_format_italics)
        #print(newword_texts)
        en_word, vi_word, pronunciation_word, type_word, out_message = analyze_line_physics_NDH(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics)
        #print(en_word)
        #print(vi_word)
        en_words.extend(en_word)
        vi_words.extend(vi_word)
        pronunciation_words.extend(pronunciation_word)
        type_words.extend(type_word)
        
        # Accumulative statistics
        if line.text.strip() == '':
            number_blank_lines += 1
        number_items_in_line[len(en_word)%11] += 1

        for k_entry in range(len(en_word)):
            if en_word[k_entry] == '' or vi_word[k_entry] == '' or pronunciation_word[k_entry] == '' or type_word[k_entry] == '':
                number_total_item_issues += 1
                out_message += '\nAttention - suspected problem with this item: ' + en_word[k_entry] + ' - ' + vi_word[k_entry] + '.'
            if en_word[k_entry] == '':
                number_en_issues += 1
                out_message += '\n     English word is missing.'
            if vi_word[k_entry] == '':
                number_vi_issues += 1
                out_message += '\n       Vietnamese word is missing.'
            if pronunciation_word[k_entry] == '':
                number_pronunciation_issues += 1
                out_message += '\n          Pronunciation word is missing.'
            if type_word[k_entry] == '':
                number_type_issues += 1
                out_message += '\n             Type word is missing.'
                
        
        out_messages.append('Line ' + str(k+1) + ': ' + out_message)
        
    # Later: output directly the result from each line to the CSV file
    table = []
    for k in range(len(en_words)):
        table.append([en_words[k], vi_words[k], pronunciation_words[k], type_words[k]])

    WriteTableCsv(exportFile, table)
    with open(logFile, 'wt') as txtfile:
        txtfile.write("Total number of lines processed: %d\n" % len(doc.paragraphs))
        txtfile.write("  with number of blank lines: %d\n" % number_blank_lines)
        txtfile.write("  \n*Number of lines failed to process (0 item): %d\n\n" % number_items_in_line[0])
        txtfile.write("  Number of lines with 1 item: %d\n" % number_items_in_line[1])
        txtfile.write("  Number of lines with 2 items: %d\n" % number_items_in_line[2])
        txtfile.write("  Number of lines with 3 items: %d\n" % number_items_in_line[3])
        txtfile.write("  Number of lines with 4 items: %d\n" % number_items_in_line[4])
        txtfile.write("  Number of lines with 5 items: %d\n" % number_items_in_line[5])
        txtfile.write("  Number of lines with 6 items: %d\n" % number_items_in_line[6])
        txtfile.write("  Number of lines with 7 items: %d\n" % number_items_in_line[7])
        txtfile.write("  Number of lines with 8 items: %d\n" % number_items_in_line[8])
        txtfile.write("  Number of lines with 9 items: %d\n" % number_items_in_line[9])
        txtfile.write("  Number of lines with 10 items: %d\n" % number_items_in_line[10])
                
        txtfile.write("\nTotal items obtained: %d\n" % len(en_words))
        txtfile.write("\n*Number of items with issues: %d\n" % number_total_item_issues)
        txtfile.write("  in which, there are:\n")
        txtfile.write("    + %d items without English word\n" % number_en_issues)
        txtfile.write("    + %d items without Vietnamese word\n" % number_vi_issues)
        txtfile.write("    + %d items without pronunciation word\n" % number_pronunciation_issues)
        txtfile.write("    + %d items without type word\n" % number_type_issues)
        txtfile.write("\n-----\n\n")
        
        for item in out_messages:
            txtfile.write("%s\n" % item)

# Main operation, when calling: python Read_OCR.py input.docx output.csv
if __name__ == "__main__":
    inputFile = str(sys.argv[1])
    if len(sys.argv)>2:
        exportFile = str(sys.argv[2])
        readocr(inputFile, exportFile)
    else:
        readocr(inputFile)
